# pages/constatator.py
from docx import Document
import re


def extrage_informatii_firma(doc):
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    company_pattern1 = r"informațiile referitoare la\s*(.*?)\s*INFORMAȚII DE IDENTIFICARE"
    company_pattern2 = r"FURNIZARE INFORMAŢII\n\n(.*?)\n"
    firma_match = re.search(company_pattern1, full_text, re.IGNORECASE | re.DOTALL)
    if firma_match:
        firma = firma_match.group(1).strip()
    else:
        firma_match = re.search(company_pattern2, full_text, re.DOTALL)
        firma = firma_match.group(1).strip() if firma_match else "N/A"
  
    nr_ordine_match = re.search(r"Număr de ordine în Registrul Comerțului:\s*([\w/]+)", full_text)
    nr_ordine = nr_ordine_match.group(1) if nr_ordine_match else "N/A"
    cui_match = re.search(r"Cod unic de înregistrare: (\d+)", full_text)
    cui = cui_match.group(1) if cui_match else "N/A"
    data_infiintarii_match = re.search(r"atribuit în data de (\d+\.\d+\.\d+)", full_text)
    data_infiintarii = data_infiintarii_match.group(1) if data_infiintarii_match else "N/A"
    address_pattern = re.compile(r"Adresă sediu social: (.*?)(?=\n)")
    address_match = re.search(address_pattern, full_text)
    adresa = address_match.group(1) if address_match else "N/A"
    main_activity_pattern = r"Activitatea principală.*?Domeniul de activitate principal:.*?\n(.*?)(?:\n|;)"
    main_activity_match = re.search(main_activity_pattern, full_text, re.DOTALL)
    main_activity = main_activity_match.group(1).strip() if main_activity_match else "N/A"
    section_pattern = re.compile(r"SEDII SECUNDARE / PUNCTE DE LUCRU(.*?)SEDII SI/SAU ACTIVITATI AUTORIZATE", re.DOTALL)
    section_match = re.search(section_pattern, full_text)
    if section_match:
        section_text = section_match.group(1)
        secondary_address_pattern = re.compile(r"Adresă: (.*?)(?=\n)", re.DOTALL)
        adrese_secundare = re.findall(secondary_address_pattern, section_text)
    else:
        adrese_secundare = ["N/A"]
    data = {
        "Denumirea firmei": firma,
        "Numărul de ordine în Registrul Comerțului": nr_ordine,
        "Codul unic de înregistrare (CUI)": cui,
        "Data înființării": data_infiintarii,
        "Adresa sediului social": adresa,
        "Activitate principală": main_activity,
        "Adresa sediul secundar": adrese_secundare
    }
    return data

def extrage_asociati_admini(doc):
    text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]  # Elimină paragrafele goale
    asociati = {}
    administratori = set()
    in_asociati_section = False
    in_persoane_imputernicite_section = False
    for i in range(len(text)):
        if "ASOCIAŢI PERSOANE FIZICE" in text[i]:
            in_asociati_section = True
        elif "REPREZENTANT acţionar/asociat/membru" in text[i]:
            in_asociati_section = False
        if in_asociati_section:
            if "Calitate: " in text[i]:
                nume = text[i - 1]
                j = i + 1
                while j < len(text) and "Cota de participare la beneficii şi pierderi: " not in text[j]:
                    j += 1
                if j < len(text):
                    cota = text[j].split(":")[1].strip()
                    asociati[nume] = cota
        if "Persoane împuternicite (PERSOANE FIZICE)" in text[i]:
            in_persoane_imputernicite_section = True
        elif "Persoane împuternicite (PERSOANE JURIDICE)" in text[i]:
            in_persoane_imputernicite_section = False
        if in_persoane_imputernicite_section and "Calitate: " in text[i]:
            nume_admin = text[i - 1]
            administratori.add(nume_admin)
    output_asociati = []
    for nume, cota in asociati.items():
        info = f"{nume} – asociat cu cota de participare la beneficii și pierderi {cota}"
        if nume in administratori:
            info += " și administrator"
        output_asociati.append(info)
    nume_administrator = ', '.join(administratori)
    return output_asociati, nume_administrator


def extrage_situatie_angajati(doc):
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    angajati_pattern_2020 = r"SITUAŢIA FINANCIARĂ PE ANUL 2020.*?(?:Numar|Număr) mediu de salari(?:aţi|ati): (\d+)"
    angajati_match_2020 = re.search(angajati_pattern_2020, full_text, re.DOTALL)
    nrang20 = angajati_match_2020.group(1) if angajati_match_2020 else "N/A"
    angajati_pattern_2021 = r"SITUAŢIA FINANCIARĂ PE ANUL 2021.*?(?:Numar|Număr) mediu de salari(?:aţi|ati): (\d+)"
    angajati_match_2021 = re.search(angajati_pattern_2021, full_text, re.DOTALL)
    nrang21 = angajati_match_2021.group(1) if angajati_match_2021 else "N/A"
    angajati_pattern_2022 = r"SITUAŢIA FINANCIARĂ PE ANUL 2022.*?(?:Numar|Număr) mediu de salari(?:aţi|ati): (\d+)"
    angajati_match_2022 = re.search(angajati_pattern_2022, full_text, re.DOTALL)
    nrang22 = angajati_match_2022.group(1) if angajati_match_2022 else "N/A"
    data_angajati = {
        "Numar mediu angajati 2020": nrang20,
        "Numar mediu angajati 2021": nrang21,
        "Numar mediu angajati 2022": nrang22,
    }
    return data_angajati

#def extrage_coduri_caen(full_text):
#    start_marker = "SEDII SI/SAU ACTIVITATI AUTORIZATE"
#    end_marker = "CONCORDAT PREVENTIV"
#    caen_section_pattern = re.compile(rf"{start_marker}(.*?){end_marker}", re.DOTALL)
#    caen_section_match = re.search(caen_section_pattern, full_text)
#    unique_caen_codes = {}

#    if caen_section_match:
#        caen_section_text = caen_section_match.group(1)
#        caen_code_pattern = re.compile(r"(\d{4}) - (.*?)\n")
#        caen_codes = re.findall(caen_code_pattern, caen_section_text)
#        for code, description in caen_codes:
#            unique_caen_codes[code] = ' '.join(description.split())
#    return list(unique_caen_codes.items())
# modificat in umra cererii din 12.feb.2024 in : 

def extrage_coduri_caen(doc):
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    start_marker = "SEDII SI/SAU ACTIVITATI AUTORIZATE"
    end_marker = "CONCORDAT PREVENTIV"
    start_index = full_text.find(start_marker) + len(start_marker)
    end_index = full_text.find(end_marker)
    relevant_section = full_text[start_index:end_index]

    pattern = r"(Sediul secundar din:.+?)(?=Sediul secundar din:|$)"
    matches = re.findall(pattern, relevant_section, re.DOTALL)

    results = []
    for match in matches:
        # Capturăm întreaga adresă a sediului până la "Tip sediu"
        sediu_info = re.search(r"(Sediul secundar din:.+?)(?=Tip sediu:)", match, re.DOTALL)
        if sediu_info:
            sediu_info = sediu_info.group(1).strip()  # Folosim group(0) pentru a include întregul match

        # Extragem activitățile la sediu și codurile CAEN
        activitati_pattern = r"Activităţi la sediu:\s*((?:\d{4} - .+?(?:\n|$))+)"
        activitati_match = re.search(activitati_pattern, match, re.DOTALL)
        if activitati_match:
            activitati_info = activitati_match.group(1).strip()
            # Eliminăm tot ce urmează după ultimul cod CAEN, inclusiv "Data certificatului constatator"
            activitati_info = re.sub(r"\nData certificatului.*$", "", activitati_info, flags=re.MULTILINE).strip()
            # Combinăm informațiile despre sediu cu activitățile la sediu
            combined_info = f"\n \n {sediu_info} \n Activităţi la sediu:\n{activitati_info}"
            results.append(combined_info)

    return results

# adaugat la cerinta din 19.02.2024 (modificari facute dupa facturare) 
def extrage_terti_principal(doc):
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    start_marker = "SEDII SI/SAU ACTIVITATI AUTORIZATE"
    end_marker = "Denumire: Punct de lucru"
    pattern = fr"(?s){start_marker}(.*?){end_marker}"

    results = []
    matches = re.findall(pattern, full_text)
    for match in matches:
        # Verificăm dacă secțiunea conține textul specificat pentru activitățile neautorizate
        if "Nu se desfăşoară activităţile prevăzute în actul constitutiv sau modificator" not in match:
            # Extragem informațiile despre tipul de activitate autorizată și codurile CAEN
            tip_activitate_pattern = r"Tip activitate autorizată: terţi\n(?:Conform declaraţiei.*?\n)?(?:Activităţi desfăşurate în afara sediului social şi a sediilor secundare \(CAEN REV\. 2\):\s*)?((?:\d{4} - .+?(?:\n|$))+)"
            tip_activitate_match = re.search(tip_activitate_pattern, match, re.DOTALL)
            if tip_activitate_match:
                tip_activitate_info = tip_activitate_match.group(1).strip()
                # Eliminăm tot ce urmează după ultimul cod CAEN, inclusiv "Data certificatului constatator"
                tip_activitate_info = re.sub(r"\nData certificatului.*$", "", tip_activitate_info, flags=re.MULTILINE).strip()
                # Combinăm informațiile despre tipul de activitate autorizată cu codurile CAEN
                combined_info = f"Tip activitate autorizată: terţi\nActivităţi desfăşurate în afara sediului social şi a sediilor secundare:\n{tip_activitate_info}"
                results.append(combined_info)

            # Extragem întreaga adresă a sediului și activitățile la sediu
            sediu_info_match = re.search(r"(Sediul (social|secundar|terţ) din:.+?)(?=Tip sediu:)", match, re.DOTALL)
            sediu_info = sediu_info_match.group(1).strip() if sediu_info_match else ""

            activitati_pattern = r"Activităţi la sediu:\s*((?:\d{4} - .+?(?:\n|$))+)"
            activitati_match = re.search(activitati_pattern, match, re.DOTALL)
            if activitati_match:
                activitati_info = activitati_match.group(1).strip()
                activitati_info = re.sub(r"\nData certificatului.*$", "", activitati_info, flags=re.MULTILINE).strip()
                # Combinăm informațiile despre sediu cu activitățile la sediu
                combined_info = f"{sediu_info}\nActivităţi la sediu:{activitati_info}"
                results.append(combined_info)
        else:
            tip_activitate_pattern = r"Tip activitate autorizată: terţi\n(?:Conform declaraţiei.*?\n)?(?:Activităţi desfăşurate în afara sediului social şi a sediilor secundare \(CAEN REV\. 2\):\s*)?((?:\d{4} - .+?(?:\n|$))+)"
            tip_activitate_match = re.search(tip_activitate_pattern, match, re.DOTALL)
            if tip_activitate_match:
                tip_activitate_info = tip_activitate_match.group(1).strip()
                # Eliminăm tot ce urmează după ultimul cod CAEN, inclusiv "Data certificatului constatator"
                tip_activitate_info = re.sub(r"\nData certificatului.*$", "", tip_activitate_info, flags=re.MULTILINE).strip()
                # Combinăm informațiile despre tipul de activitate autorizată cu codurile CAEN
                combined_info = f"Tip activitate autorizată: terţi\nActivităţi desfăşurate în afara sediului social şi a sediilor secundare:\n{tip_activitate_info}"
                results.append(combined_info)


    return results
