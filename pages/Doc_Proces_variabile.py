import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def ajusteaza_alinierea(document):
    for paragraph in document.paragraphs:
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def main():
    st.header("Ajustare Aliniere Document")

    uploaded_file = st.file_uploader("Încărcați documentul DOCX aici:", type="docx")
    if uploaded_file is not None:
        document = Document(uploaded_file)
        ajusteaza_alinierea(document)
        
        # Salvăm documentul modificat într-un buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        st.download_button(label="Descarcă Documentul Modificat",
                           data=buffer,
                           file_name="document_modificat.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()
