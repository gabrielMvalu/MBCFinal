# pages/Completare_Doc.py
import streamlit as st
import pandas as pd
import re
import math
from docx import Document
from constatator import extrage_informatii_firma, extrage_asociati_admini, extrage_situatie_angajati, extrage_coduri_caen, extrage_terti_principal
from datesolicitate import extrage_date_solicitate, extrage_date_suplimentare
from bilantsianaliza import extrage_date_bilant, extrage_date_contpp, extrage_indicatori_financiari
from serviciisiutilaje import extrage_pozitii, coreleaza_date
from cheltuieliBugetVar import extrage_cheltuieli_buget

st.set_page_config(layout="wide")

st.header(':blue[Procese pentru inlocuire "Variabile" si completare "Machete PA/CF..."]', divider='rainbow')

if 'downloaded' not in st.session_state:
    st.session_state['downloaded'] = False

nr_CAEN = None

document_succes = False  
document2_succes = False
document3_succes = False


col1, col2 = st.columns(2)

with col1:
    uploaded_doc1 = st.file_uploader("Încărcați fișierul Date Solicitate", type=["xlsx"], key="dateSolicitate")
    if uploaded_doc1 is not None:
        df = pd.read_excel(uploaded_doc1)
        solicitate_data = extrage_date_solicitate(df)

        firma = solicitate_data.get('Denumirea firmei SRL', 'N/A')
        nr_CAEN = solicitate_data.get('Doar nr CAEN','N/A')
        st.session_state.codCAEN = nr_CAEN
        st.success(f"Primul pas, pentru: {firma}, completat.")
        document_succes = True
        descriere_u_r = solicitate_data.get('D u reciclare', 'N/A')
    


with col2:
    if document_succes:
        uploaded_doc2 = st.file_uploader(f"Încărcați 'RAPORT INTEROGARE' al: {firma}", type=["docx"], key="RaportInterogare")
        if uploaded_doc2 is not None:
            constatator_doc = Document(uploaded_doc2)
    
            informatii_firma = extrage_informatii_firma(constatator_doc)
            asociati_info, administratori_info = extrage_asociati_admini(constatator_doc)
            situatie_angajati = extrage_situatie_angajati(constatator_doc)
            # full_text_constatator = "\n".join([p.text for p in constatator_doc.paragraphs])
            sedii_si_activitati = extrage_coduri_caen(constatator_doc)
            terti_si_principal = extrage_terti_principal(constatator_doc) # adaugat dupa facturare (la cerinte noi) 19.02
            srl = informatii_firma.get('Denumirea firmei', 'N/A')
            
            # modificari adaugare variabile noi 16 feb 2024
       
            # adaugata 16.feb 2024 pt nr mediu angajati pe 2022 din constatator       
            # Încercăm să extragem și să convertim valoarea pentru anul 2022
            nr_ang_brut_22 = situatie_angajati.get('Numar mediu angajati 2022', 'N/A')
            try:
                #convertim valoarea într-un întreg
                nr_angajati_22 = int(nr_ang_brut_22)
            except ValueError:
                nr_angajati_22 = None  
            # La acest punct, `nr_angajati_22` este fie un întreg, fie `None`  
            
            #modificat cnf noilor cerinte din:12.feb.2024 
            #coduri_caen = extrage_coduri_caen(full_text_constatator)
                      
            #def curata_duplicate_coduri_caen(coduri_caen):
            #    coduri_unice = {}
            #    for cod, descriere in coduri_caen:
            #        coduri_unice[cod] = descriere
            #    return list(coduri_unice.items())
    
            #coduri_caen_curatate = curata_duplicate_coduri_caen(coduri_caen)
            adrese_secundare_text = '\n'.join(informatii_firma.get('Adresa sediul secundar', [])) if informatii_firma.get('Adresa sediul secundar', []) else "N/A"
            asociati_text = '\n'.join(asociati_info) if asociati_info else "N/A"
            administratori_text = administratori_info if administratori_info else "N/A"
            sedii_si_activitati_text = '\n'.join(sedii_si_activitati) if sedii_si_activitati  else "N/A"           
            #coduri_caen_text = '\n'.join([f"{cod} - {descriere}" for cod, descriere in coduri_caen_curatate]) if coduri_caen_curatate else "N/A"    
            terti_si_principal_text = '\n'.join(terti_si_principal) if terti_si_principal  else "N/A"
            st.info(f"Prelucrarea 'Rapor Interogare' al {srl}, este completa.")
            document2_succes = True        

            # modificari pentru avertizarea in cazul documentelor ce nu sunt standardizate pt 3 ani in special 20,21,22
             
            nrang20 = situatie_angajati.get("Numar mediu angajati 2020")
            nrang21 = situatie_angajati.get("Numar mediu angajati 2021")
            nrang22 = situatie_angajati.get("Numar mediu angajati 2022")
            # Verificarea dacă oricare dintre ani nu are date extrase corespunzător
            if nrang20 == "N/A" or nrang21 == "N/A" or nrang22 == "N/A":
                st.info(f"Documentul nu este standardizat {nrang20} {nrang21} {nrang22}, continuarea executării poate duce la generarea de date eronate. Asigurați-vă că 'Raportul de interogare' corespunde cerințelor!")
    
            #  Afișarea datelor în format JSON
 #           st.json({
 #               "Date Generale": informatii_firma,
 #               "Informații Detaliate": {"Asociați": asociati_info, "Administratori": administratori_info},
 #               "Situație Angajati": situatie_angajati,
 #               "Coduri CAEN": sedii_si_activitati
 #           }) 
    else:
        st.warning("Prima dată, încărcați și procesați, 'date solicitate.xlsx'.")



col3, col4 = st.columns(2)

with col3:
    if document2_succes:
        uploaded_doc3 = st.file_uploader("Încărcați Anexa 3 Macheta Financiară", type=["xlsx"], key="AnalizaMacheta")
        if uploaded_doc3 is not None:
            df_bilant = pd.read_excel(uploaded_doc3, sheet_name='1-Bilant')
            df_contpp = pd.read_excel(uploaded_doc3, sheet_name='2-ContPP')
            df_analiza_fin = pd.read_excel(uploaded_doc3, sheet_name='1D-Analiza_fin_indicatori')    
            df_financiar = pd.read_excel(uploaded_doc3, sheet_name='P. FINANCIAR')
            
            date_financiare = extrage_pozitii(df_financiar)
           

            if nr_CAEN != 'N/A' and date_financiare:
               # rezultate_corelate, rezultate_corelate1, rezultate_corelate2 = coreleaza_date(date_financiare)
                rezultate_corelate, rezultate_corelate2 = coreleaza_date(date_financiare)
                
                rezultate_text = '\n'.join([rezultat for _, _, rezultat in rezultate_corelate])
               # cheltuieli_text = '\n'.join([rezultat for _, _, rezultat in rezultate_corelate1]) - Eliminat datorita schimbarii variabilei cheltuieli_buget
                df_cheltuieli = extrage_cheltuieli_buget(df_financiar) #adaugat dupa eliminarea rezultate_corelate1
                cheltuieli_buget_text = ''.join([f"{row.iloc[0]} - {row.iloc[1]} buc. \n" for _, row in df_cheltuieli.iterrows()])  #adaugat dupa eliminarea rezultate_corelate1
                
                cantitati_corelate = [pd.to_numeric(item[1], errors='coerce') for item in rezultate_corelate]
                cantitati_corelate = [0 if pd.isna(x) else x for x in cantitati_corelate]
                numar_total_utilaje = sum(cantitati_corelate)
                rezultate_corelate, rezultate_corelate2 = coreleaza_date(date_financiare)
                rezultate2_text = '\n'.join([f"{nume} - {descriere}" for nume, _, descriere in rezultate_corelate2])

                nrutlocm_temp = df_financiar.iloc[4, 21] # adaugat in 12 feb pt modificarile facute legat de extragerea nr total de noi angajati conf proiect
                
                procentCrestereNrLocMunca = df_financiar.iloc[3, 22] / 100 # adaugat cf cerinte din 16 feb
                procentCrestereAngajati = f"{procentCrestereNrLocMunca:.2%}"
                
                procentCAtemp = df_financiar.iloc[32, 18] / 100
                procentCA = f"{procentCAtemp:.2%}"
                
                valoareCheltuieliNeeligibile = df_financiar.iloc[24, 7]
                procentNePerValoareTPtemp = df_financiar.iloc[24, 7] / df_financiar.iloc[24, 2]
                procentNePerValoareTP = f"{procentNePerValoareTPtemp:.2%}"
                
                totalEligibil = df_financiar.iloc[24, 6]
                valAFN = df_financiar.iloc[24, 8]    #end partial 16

                # modific pt funct 16
                ca22 = df_contpp.iloc[4, 3]
                raportCA22AFNtemp = ca22 / valAFN    # Raportul dintre cuantumul finanțării solicitate si cifra de afaceri înregistrată în anul fiscal anterior lansării apelului de proiecte                
                raportCA22AFN = f"{raportCA22AFNtemp: .2f}"   #end modfiic funct 16
               

                # mai jos pregatire pt modificari 16 feb 
                try:
                    nrLocMuncaNoi = int(nrutlocm_temp)
                except ValueError:
                    nrLocMuncaNoi = None  

                # mai jos : modificari 16 feb la cerinta: calcul variabile noi in cod
                if nr_angajati_22 is not None and nrLocMuncaNoi is not None:
                     try:
                        totalAngajati = nrLocMuncaNoi + nr_angajati_22
                        nrLocMunca30 = math.ceil(nrLocMuncaNoi * 0.30) 
                        nrLocMunca20 = math.ceil(nrLocMuncaNoi * 0.20) 
                        procentTotalSi30 = nrLocMunca30 + nr_angajati_22
                        procentTotalSi20 = nrLocMunca20 + nr_angajati_22
                     except ValueError:
                         st.error("Una dintre valorile pentru calculul totalului angajaților nu este un număr valid.")
                        
                else:
                    st.error("Una dintre valorile necesare pentru calculul totalului angajaților lipsește sau este None.")     #modific 16 end aici

                
            if nrLocMuncaNoi < numar_total_utilaje:
                nrutverificat1 = "De aceea, nu este necesar să angajam atât de mulți muncitori câte utilaje avem, ci să ne adaptam forța de muncă la nevoile specifice ale proiectelor pe care le vom executa."
                nrutverificat2 = "Mai mult decat atat, daca va fi cazul, pentru operarea utilajelor prevazute in cadrul proiectului  va fi folosit o parte din personal existent prin relocare si utilizarea exclusiva in cadrul acestui proiect si se va angaja si personal nou, calificat/necalificat, in functie de necesitatile existente la un moment dat pentru aceasta activitate."
            else:
                nrutverificat1 = "."
                nrutverificat2 = "."            
            

            capital_propriu = extrage_date_bilant(df_bilant)
            cifra_venit_rezultat = extrage_date_contpp(df_contpp)
            rata_rent_grad = extrage_indicatori_financiari(df_analiza_fin)

            st.success(f"Analiza Financiara prelucrata cu succes. Va rugam Adaugati Macheta PA si completati procesul.")
            document3_succes = True
    else:
        st.warning("Vă rugăm să încărcați și să procesați 'Date Solicitate', apoi 'Raport Interogare'.")


with col4:
    if document3_succes and 'downloaded' in st.session_state and not st.session_state['downloaded']:
        uploaded_template = st.file_uploader("Încărcați MACHETA pt procesarea finala.", type=["docx"], key="MachetaPA")
        if uploaded_template is not None:
            template_doc = Document(uploaded_template)
            st.toast('A inceput procesarea Planului de afaceri', icon='⭐')
            # st.info(f"Procesare Finalizata. Pentru descarcarea documentului completat", icon="⬇️")
            st.info(f"Procesare Finalizata. Asteptati Butonul pentru descarcarea documentului completat ")
            
            # Preia valoarea pentru 'Utilaj cu tocător' - Eliminat in urma schimbarii variabilei ReciclareaMaterialelor.
        #    utilaj_cu_tocator_pt_inlocuire = solicitate_data.get('Utilaj cu tocător', 'N/A')
            nr_clasare_notificare_pt_inlocuire = solicitate_data.get('Număr clasare notificare', 'N/A')
            
            # Preia și prelucrează valoarea pentru 'Procesul de reciclare a materialelor'
        #    reciclareaMaterialelor_temp = solicitate_data.get('Procesul de reciclare a materialelor', 'N/A')
        #    reciclareaMaterialelor_complet = reciclareaMaterialelor_temp.replace("#utilaj_cu_tocator", utilaj_cu_tocator_pt_inlocuire)

            # Preia și prelucrează valoarea pentru 'Nr clasare notificare'
            cDNSH_temp = solicitate_data.get('Detalii DNSH - C', 'N/A')
            cDNSH_complet = cDNSH_temp.replace("#nr_clasare_notificare", nr_clasare_notificare_pt_inlocuire)

            dDNSH_text = solicitate_data.get('Detalii DNSH - D', 'N/A')
            
            
            placeholders = {
                "#SRL": str(informatii_firma.get('Denumirea firmei', 'N/A')),
                "#CUI": str(informatii_firma.get('Codul unic de înregistrare (CUI)', 'N/A')),
                "#Nr_inmatriculare": str(informatii_firma.get('Numărul de ordine în Registrul Comerțului', 'N/A')),
                "#data_infiintare": str(informatii_firma.get('Data înființării', 'N/A')),
                "#Adresa_sediu": str(informatii_firma.get('Adresa sediului social', 'N/A')),
                "#Adresa_pct_lucru": str(adrese_secundare_text),
                "#Asociati": str(asociati_text),
                "#Administrator": str(administratori_text),
                "#activitatePrincipala": str(informatii_firma.get('Activitate principală', 'N/A')),

                
                "#CAENautorizate": str(sedii_si_activitati_text), #modificar la variabila dupa 12.02 ( indicatii eronate initial date, apoi refacute, apoi iar eronata XXX
                "#tertiPrincipal": str(terti_si_principal_text), #adauat dupa facturarea pe 19.02
                
                "#categ_intreprindere": str(solicitate_data.get('Categorie întreprindere', 'N/A')),
                "#Firme_legate": str(solicitate_data.get('Firme legate', 'N/A')),
                "#Tip_investitie": str(solicitate_data.get('Tipul investiției', 'N/A')),
                "#activitate": str(solicitate_data.get('Activitate', 'N/A')),
                "#CAEN": str(solicitate_data.get('Cod CAEN', 'N/A')),
                
            #    "#nr_locuri_munca_noi": str(solicitate_data.get('Număr locuri de muncă noi', 'N/A')),
                "#nr_locuri_munca_noi": str(nrLocMuncaNoi), #modificat la cerere pt a nu se mai prelua valoarea din date_financiare
                
                
                "#Judet": str(solicitate_data.get('Județ', 'N/A')),
        
                "#Utilaj": str(rezultate_text),
               # "#cheltuieli_proiect_din_buget_excel": str(cheltuieli_text),
                "#cheltuieli_proiect_din_buget_excel": str(cheltuieli_buget_text), #adaugat in urma modif din 14.feb.2024 pt modificarea variabilei
                
                "#DescriereUtilaje" : str(rezultate2_text),
                "#nr_utilaje": str(numar_total_utilaje),
                
                "#utilaj_dizabilitati": str(solicitate_data.get('Utilaj pentru persoane cu dizabilități', 'N/A')),
                
          #      "#utilaj_cu_tocator": str(solicitate_data.get('Utilaj cu tocător', 'N/A')),
                "#adresa_loc_implementare": str(solicitate_data.get('Adresa locației de implementare', 'N/A')),
                "#nrClasareNotificare": str(solicitate_data.get('Număr clasare notificare', 'N/A')),
                "#clientiActuali": str(solicitate_data.get('Clienți actuali', 'N/A')),
                "#furnizori": str(solicitate_data.get('Furnizori', 'N/A')),
                "#tip_activitate": str(solicitate_data.get('Tip activitate', 'N/A')),
                "#ISO": str(solicitate_data.get('Certificări ISO', 'N/A')),
                "#CurentaActivitate": str(solicitate_data.get('Curenta Activitate', 'N/A')),
                "#dotari_activitate_curenta": str(solicitate_data.get('Dotări pentru activitatea curentă', 'N/A')),
                "#info_ctr_implementare": str(solicitate_data.get('Informații despre contractul de implementare', 'N/A')),
                "#zonele_vizate_prioritar": str(solicitate_data.get('Zonele vizate prioritare', 'N/A')),
                "#utilaj_ghidare": str(solicitate_data.get('Utilaj de ghidare', 'N/A')),
                
                "#legaturi": str(solicitate_data.get('Legaturi', 'N/A')),
                "#rude": str(solicitate_data.get('Rude', 'N/A')),
                
                "#detaliere_CA": str(solicitate_data.get('Detalierea CA', 'N/A')),
                "#legate_detaliere_CA": str(solicitate_data.get('Detaliere CA legate', 'N/A')),
                "#posturiRevisal": str(solicitate_data.get('Posturi Revisal', 'N/A')),
                
                "#concluzie_CA": str(solicitate_data.get('Concluzie cifra de afaceri', 'N/A')),
                
                "#caracteristici_tehnice": str(solicitate_data.get('Caracteristici tehnice utilaje', 'N/A')),
                "#flux_tehnologic": str(solicitate_data.get('Fluxul tehnologic', 'N/A')),
                "#utilajeDNSH": str(solicitate_data.get('DNSH pentru utilaje', 'N/A')),
                
                "#descriere_utilaj_ghidare": str(solicitate_data.get('Descrierea utilaj ghidare', 'N/A')),               
                "#descriereUtilajReciclare": str(descriere_u_r),
                
                "#contributia_proiectului_la_TJ": str(solicitate_data.get('Contribuția proiectului la tranziția justă', 'N/A')),
                "#strategii_materiale": str(solicitate_data.get('Strategii materiale', 'N/A')),
                "#strategii_reciclate": str(solicitate_data.get('Strategii materiale reciclate', 'N/A')),
                "#activitate": str(solicitate_data.get('Activitate specifică', 'N/A')),
                "#descriere_utilaj_reciclare": str(solicitate_data.get('Descriere utilaj de reciclare', 'N/A')),
                "#lucrari_inovatie": str(solicitate_data.get('Inovații în lucrări', 'N/A')),
                "#lucrari_caen": str(solicitate_data.get('Lucrări conform codurilor CAEN', 'N/A')),
                "#aDNSH": str(solicitate_data.get('Detalii DNSH - A', 'N/A')),
                "#cDNSH": str(cDNSH_complet),
                "#dDNSH": str(dDNSH_text),
                "#materiale_locale": str(solicitate_data.get('Utilizarea materialelor locale', 'N/A')),
                "#PregatireaTeren": str(solicitate_data.get('Pregătirea terenului pentru lucrări', 'N/A')),
                "#ReciclareaMaterialelor": str(solicitate_data.get('Procesul de reciclare a materialelor', 'N/A')),
                "#clientiFirma": str(solicitate_data.get('Clienți principali ai firmei', 'N/A')),
                "#DacaTipInvest": str(solicitate_data.get('Tipul investitiei', 'N/A')),
                "#crearea": str(solicitate_data.get('Crearea de noi oportunități', 'N/A')),
                "#CompletDiversificare": str(solicitate_data.get('Diversificarea activităților firmei', 'N/A')),
                "#CompletExtinderea": str(solicitate_data.get('Extinderea capacității firmei', 'N/A')),
                "#CrestCreare": str(solicitate_data.get('Creșterea sau crearea de noi surse de venit', 'N/A')),
                "#CreareActivVizata": str(solicitate_data.get('Crearea de activități în domeniul vizat', 'N/A')),
                "#DezavantajeConcurentiale": str(solicitate_data.get('Identificarea dezavantajelor concurențiale', 'N/A')),
                
                "#30nrLocMunca": str(nrLocMunca30), #modificari aduse in 16 feb 2024 pt 30 si pt 20, 30Procent si 20 Procent
                "#20NrLocMunca": str(nrLocMunca20),
                
                "#30%total": str(procentTotalSi30), #var noi adaugate post facturare pt CF referitor la anagajati
                "#20%total": str(procentTotalSi20),
                "#%crestereang": str(procentCrestereAngajati),
                "#totalangajati": str(totalAngajati), #end var ang noi
        
                "#valoare_cheltuieli_neeligibile": str(valoareCheltuieliNeeligibile), #var noi adaugate post fact pt CF referitor la chelt neelig
                "#procent_neeligibil_din_valoare_totala_proiect": str(procentNePerValoareTP),    
                "#%crestere_CA": str(procentCA),
                "#total_eligibil": str(totalEligibil),  
                "#AFN": str(valAFN),
                "#raportCA_AFN": str(raportCA22AFN),     # end modif 16.02 / eng variabile ref chelt
                
                "#zoneDN": str(solicitate_data.get('Zone vizate Prioritar', 'N/A')),
                "#Iso14001": str(solicitate_data.get('Daca are sau nu iso14001', 'N/A')),
                "#descriere_serviciu": str(solicitate_data.get('Descriere serviciului', 'N/A')),
                "#piata_tinta": str(solicitate_data.get('Piata tinta', 'N/A')),
                "#clientiFirma": str(solicitate_data.get('Clienți principali ai firmei', 'N/A')),

                #necesaitau modific in urma dec 12.02
                
                "#NlmU1": str(nrutverificat1),
                "#NlmU2": str(nrutverificat2),                             

                   #end modif pt 12 feb
                            
                "#NAM20": str(situatie_angajati.get('Numar mediu angajati 2020', 'N/A')),
                "#NAM21": str(situatie_angajati.get('Numar mediu angajati 2021', 'N/A')),
                "#NAM22": str(situatie_angajati.get('Numar mediu angajati 2022', 'N/A')),   
                "#CPA20": str(capital_propriu.get('Capitalul propriu al actionarilor 2020', 'N/A')), 
                "#CPA21": str(capital_propriu.get('Capitalul propriu al actionarilor 2021', 'N/A')),
                "#CPA22": str(capital_propriu.get('Capitalul propriu al actionarilor 2022', 'N/A')),
                "#CA20": str(cifra_venit_rezultat.get('Cifra de afaceri 2020', 'N/A')),
                "#CA21": str(cifra_venit_rezultat.get('Cifra de afaceri 2021', 'N/A')),
                "#CA22": str(cifra_venit_rezultat.get('Cifra de afaceri 2022', 'N/A')),
                "#VT20": str(cifra_venit_rezultat.get('Venituri totale 2020', 'N/A')),
                "#VT21": str(cifra_venit_rezultat.get('Venituri totale 2021', 'N/A')),
                "#VT22": str(cifra_venit_rezultat.get('Venituri totale 2022', 'N/A')),     
                "#REX20": str(cifra_venit_rezultat.get('Rezultat al exercitiului 2020', 'N/A')),
                "#REX21": str(cifra_venit_rezultat.get('Rezultat al exercitiului 2021', 'N/A')), 
                "#REX22": str(cifra_venit_rezultat.get('Rezultat al exercitiului 2022', 'N/A')),
                "#MAXCA": str(cifra_venit_rezultat.get('Anul cu cea mai mare cifra de afaceri', 'N/A')),
                "#RSG20": str(rata_rent_grad.get('Rata solvabilitatii generale 2020', 'N/A')),
                "#RSG21": str(rata_rent_grad.get('Rata solvabilitatii generale 2021', 'N/A')), 
                "#RSG22": str(rata_rent_grad.get('Rata solvabilitatii generale 2022', 'N/A')),
                "#GITS20": str(rata_rent_grad.get('Gradul de indatorare pe termen scurt 2020', 'N/A')),
                "#GITS21": str(rata_rent_grad.get('Gradul de indatorare pe termen scurt 2021', 'N/A')),
                "#GITS22": str(rata_rent_grad.get('Gradul de indatorare pe termen scurt 2022', 'N/A')),
                "#ROA20": str(rata_rent_grad.get('Rentabilitatea activelor (ROA) 2020', 'N/A')),
                "#ROA21": str(rata_rent_grad.get('Rentabilitatea activelor (ROA) 2021', 'N/A')),
                "#ROA22": str(rata_rent_grad.get('Rentabilitatea activelor (ROA) 2022', 'N/A')),
                "#ROE20": str(rata_rent_grad.get('Rentabilitatea capitalului propriu (ROE) 2020', 'N/A')),
                "#ROE21": str(rata_rent_grad.get('Rentabilitatea capitalului propriu (ROE) 2021', 'N/A')),
                "#ROE22": str(rata_rent_grad.get('Rentabilitatea capitalului propriu (ROE) 2022', 'N/A')),
                #schimbari 16 feb pt CF
                "#RPE22": str(rata_rent_grad.get('Rata profitului din exploatare RPE22', 'N/A')),
                "#REF22": str(rata_rent_grad.get('Rentabilitatea capitalului propriu (ROE) 2022', 'N/A')), # end here sch 16 CF

            }
        
            def inlocuieste_in_tabele(tabele, placeholders):
                for tabel in tabele:
                    for row in tabel.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    for placeholder, value in placeholders.items():
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, value)
                            # Verifică dacă există tabele încastrate în celulă și aplică funcția recursiv
                            if cell.tables:
                                inlocuieste_in_tabele(cell.tables, placeholders)
        
            inlocuieste_in_tabele(template_doc.tables, placeholders)
        
            for paragraph in template_doc.paragraphs:
                for run in paragraph.runs:
                    for placeholder, value in placeholders.items():
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
        
            template_doc.save("plan_afaceri_completat.docx")

            # st.info(f"Procesare Finalizata. Asteptati Butonul pentru descarcarea documentului completat ")
            
            with open("plan_afaceri_completat.docx", "rb") as file:
                st.download_button(label="Descarcă Documentul Completat", data=file, file_name="document_modificat.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.session_state['downloaded'] = True
                         

    else:
        st.warning("Vă rugăm să încărcați și să procesați documentele din primele coloane.")


